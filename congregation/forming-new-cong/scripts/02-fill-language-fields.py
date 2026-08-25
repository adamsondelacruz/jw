#!/usr/bin/env python3
"""Create the evidence-backed working copy of Daniel's language-fields form.

The received DOCX is immutable. This script writes a separate working copy and
uses blue text for entered values so that additions are visually distinct.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DEFAULT_VALUES = (
    "Ashburton Tagalog Group",
    "5",
    "2",
    "31",
    "47",
    "TBC",
)


def set_cell(cell, value: str, *, centered: bool) -> None:
    paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = paragraph.add_run(value)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)


def build(source: Path, output: Path) -> None:
    source = source.resolve()
    output = output.resolve()
    if source == output:
        raise SystemExit("Refusing to overwrite the received source document.")
    if not source.is_file():
        raise SystemExit(f"Source document not found: {source}")

    document = Document(source)
    if len(document.tables) != 2:
        raise SystemExit("Unexpected form layout: expected exactly two tables.")
    table = document.tables[1]
    if len(table.rows) < 2 or len(table.columns) != 6:
        raise SystemExit("Unexpected groups/pregroups table layout.")

    for index, value in enumerate(DEFAULT_VALUES):
        set_cell(table.rows[1].cells[index], value, centered=index > 0)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build(args.source, args.output)


if __name__ == "__main__":
    main()
